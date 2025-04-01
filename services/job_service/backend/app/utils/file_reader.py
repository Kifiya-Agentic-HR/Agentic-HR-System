import os
import logging
import requests
import tempfile
from pdfminer.high_level import extract_text
from docx import Document
import fitz  # PyMuPDF
import pytesseract
from PIL import Image

logger = logging.getLogger(__name__)

def extract_text_from_file(file_path_or_url: str) -> str:
    """
    Handles both local files and URLs for text extraction
    Supports PDF (including image-based), DOCX, DOC, and TXT files
    """
    try:
        # Check if input is a local file path
        if os.path.exists(file_path_or_url):
            logger.info(f"Processing local file: {file_path_or_url}")
            
            # Handle different file types
            if file_path_or_url.lower().endswith('.pdf'):
                return extract_text_from_pdf(file_path_or_url)
            elif file_path_or_url.lower().endswith(('.docx', '.doc')):
                return extract_text_from_docx(file_path_or_url)
            elif file_path_or_url.lower().endswith('.txt'):
                return extract_text_from_txt(file_path_or_url)
            else:
                raise ValueError(f"Unsupported file format: {file_path_or_url}")
        
        # Handle URL case
        logger.info(f"Processing URL: {file_path_or_url}")
        response = requests.get(file_path_or_url)
        response.raise_for_status()
        
        # Handle different content types
        content_type = response.headers.get('Content-Type', '')
        if 'application/pdf' in content_type:
            return extract_text_from_pdf_bytes(response.content)
        elif 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' in content_type:
            return extract_text_from_docx_bytes(response.content)
        elif 'text/plain' in content_type:
            return response.text
        else:
            raise ValueError(f"Unsupported content type from URL: {content_type}")

    except Exception as e:
        logger.error(f"Error extracting text from {file_path_or_url}: {str(e)}")
        raise

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF with OCR fallback"""
    try:
        # First try standard text extraction
        text = extract_text(file_path)
        
        # Check if text extraction might have failed (image-based PDF)
        if len(text.strip()) < 50:  # Adjust threshold as needed
            logger.info("Low text count, attempting OCR")
            return extract_text_from_image_pdf(file_path)
        return text
    except Exception as e:
        logger.error(f"PDF extraction error: {str(e)}")
        raise

def extract_text_from_image_pdf(file_path: str) -> str:
    """Extract text from image-based PDF using OCR"""
    text = []
    try:
        doc = fitz.open(file_path)
        for page_num, page in enumerate(doc):
            # Render page as an image
            pix = page.get_pixmap(dpi=300)  # Increase DPI for better OCR accuracy
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            # Perform OCR
            page_text = pytesseract.image_to_string(img)
            text.append(f"--- PAGE {page_num + 1} ---\n{page_text}")
            
            # Clean up resources
            del pix
            
        return '\n'.join(text)
    except Exception as e:
        logger.error(f"OCR extraction error: {str(e)}")
        raise
    finally:
        if 'doc' in locals():
            doc.close()

def extract_text_from_pdf_bytes(content: bytes) -> str:
    """Extract text from PDF bytes with OCR fallback"""
    try:
        with tempfile.NamedTemporaryFile(delete=True, suffix='.pdf') as tmp:
            tmp.write(content)
            tmp.seek(0)
            
            # First try standard extraction
            text = extract_text(tmp.name)
            if len(text.strip()) < 50:
                logger.info("Low text count, attempting OCR for byte content")
                return extract_text_from_image_pdf(tmp.name)
            return text
    except Exception as e:
        logger.error(f"PDF bytes extraction error: {str(e)}")
        raise

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from local DOCX file"""
    try:
        doc = Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])
    except Exception as e:
        logger.error(f"DOCX extraction error: {str(e)}")
        raise

def extract_text_from_txt(file_path: str) -> str:
    """Extract text from local TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        logger.error(f"TXT extraction error: {str(e)}")
        raise

def extract_text_from_docx_bytes(content: bytes) -> str:
    """Extract text from DOCX bytes"""
    try:
        with tempfile.NamedTemporaryFile(delete=True, suffix='.docx') as tmp:
            tmp.write(content)
            tmp.seek(0)
            doc = Document(tmp.name)
            return '\n'.join([para.text for para in doc.paragraphs])
    except Exception as e:
        logger.error(f"DOCX bytes extraction error: {str(e)}")
        raise